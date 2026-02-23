import os
from playwright.async_api import async_playwright
from langchain_core.tools import tool
import docx
import pdfplumber
import subprocess
from io import BytesIO

@tool
def extract_resume_text(file_path: str):
    """
        Extracts the text content from a CV file in PDF, DOCX, or DOC format. The CV should be inside the same folder that this notebook is in.

        Args:
            file_path (str): The local file path to the CV document. 
        
        Returns:
            str: The extracted plain text from the CV, or an error message if the format is unsupported or cannot be read.
    """
    extension = os.path.splitext(file_path)[-1].lower()

    if '.docx' in extension:
        try:
            # document = docx.Document(file_path)
            document = docx.Document(file_path)
            text = [para.text for para in document.paragraphs]
            return '\n'.join(text)
        except Exception as e:
            print("Error processing .docx file; ", e)
    
    elif '.pdf' in extension:
        try:
            text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_content = page.extract_text()
                    if page_content:
                        text.append(page_content)
            return '\n'.join(text)
        except Exception as e:
            print('Error processing the PDF file :', e)

    elif '.doc' in extension:
        try:
            temp_docx = file_path + '.temp.docx'
            subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--ootdir', os.path.dirname(file_path), file_path], check=True)
            doc = docx.Document(os.path.splitext(file_path)[0] + ".docx")
            text = [para for para in doc.paragraphs]

            os.remove(os.path.splitext(file_path)[0] + ".docx")
            return '\n'.join(text)
        except Exception as e:
            print("Error processing .doc file: ", e)
    else:
        return "Unsupported file type. Please upload your CV in .pdf, .docx, or .doc format."